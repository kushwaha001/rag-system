from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
import os
from datetime import date


def generate_docx(paper_data: dict, topic: str, difficulty: str) -> str:
    doc = Document()

    # ── PAGE MARGINS ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    today = date.today().strftime("%B %d, %Y")
    total_q = (
        len(paper_data.get("mcq", []))
        + len(paper_data.get("short", []))
        + len(paper_data.get("long", []))
    )

    # ── HEADER ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("QUESTION PAPER")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Subject: {topic}").bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Difficulty: {difficulty}    |    Total Questions: {total_q}    |    Date: {today}")
    meta.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("─" * 70)

    # ── QUESTIONS ──
    qno = 1

    if paper_data.get("mcq"):
        h = doc.add_heading("Section A — Multiple Choice Questions", level=1)
        doc.add_paragraph("")
        for q in paper_data["mcq"]:
            p = doc.add_paragraph()
            r = p.add_run(f"Q{qno}. ")
            r.bold = True
            p.add_run(q.get("question", ""))
            for key, opt in (q.get("options") or {}).items():
                doc.add_paragraph(f"        {key}) {opt}")
            doc.add_paragraph("")
            qno += 1

    if paper_data.get("short"):
        doc.add_heading("Section B — Short Answer Questions", level=1)
        doc.add_paragraph("")
        for q in paper_data["short"]:
            p = doc.add_paragraph()
            p.add_run(f"Q{qno}. ").bold = True
            p.add_run(q.get("question", ""))
            doc.add_paragraph("")
            qno += 1

    if paper_data.get("long"):
        doc.add_heading("Section C — Long Answer Questions", level=1)
        doc.add_paragraph("")
        for q in paper_data["long"]:
            p = doc.add_paragraph()
            p.add_run(f"Q{qno}. ").bold = True
            p.add_run(q.get("question", ""))
            doc.add_paragraph("")
            qno += 1

    # ── PAGE BREAK ──
    doc.add_page_break()

    # ── ANSWER KEY ──
    ak_title = doc.add_paragraph()
    ak_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ak_title.add_run("ANSWER KEY")
    r.bold = True
    r.font.size = Pt(18)

    ak_sub = doc.add_paragraph()
    ak_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ak_sub.add_run(f"Subject: {topic}    |    {difficulty}").bold = True

    doc.add_paragraph("─" * 70)
    doc.add_paragraph("")

    qno = 1

    if paper_data.get("mcq"):
        doc.add_heading("Section A — MCQ Answers", level=1)
        for q in paper_data["mcq"]:
            p = doc.add_paragraph()
            p.add_run(f"Q{qno}. ").bold = True
            ans_letter = str(q.get("answer", "—"))
            ans_text = (q.get("options") or {}).get(ans_letter, "")
            ans_display = f"{ans_letter}) {ans_text}" if ans_text else ans_letter
            p.add_run(ans_display)
            qno += 1
        doc.add_paragraph("")

    if paper_data.get("short"):
        doc.add_heading("Section B — Short Answer Key", level=1)
        for q in paper_data["short"]:
            p = doc.add_paragraph()
            p.add_run(f"Q{qno}. ").bold = True
            p.add_run(str(q.get("answer", "—")))
            qno += 1
        doc.add_paragraph("")

    if paper_data.get("long"):
        doc.add_heading("Section C — Long Answer Key", level=1)
        for q in paper_data["long"]:
            p = doc.add_paragraph()
            p.add_run(f"Q{qno}. ").bold = True
            p.add_run(str(q.get("answer", "—")))
            qno += 1

    filename = f"question_paper_{uuid.uuid4().hex}.docx"
    filepath = os.path.join("/tmp", filename)
    doc.save(filepath)
    return filepath
